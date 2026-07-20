from fastapi import APIRouter, Depends, Request, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.database.session import get_db
from app.schemas.chat import (
    MessageCreate, 
    MessageResponse, 
    ConversationResponse, 
    ConversationRenameRequest,
    FeedbackRequest,
    BookmarkRequest
)
from app.services.chat_service import ChatService
from app.core.security import get_current_user
from app.models.user import User
from typing import List, Dict, Any, Optional

from app.repositories.chat_repo import ChatRepository
from app.repositories.stats_repo import StatsRepository

def get_chat_service(db: Session = Depends(get_db)) -> ChatService:
    return ChatService(
        db=db,
        chat_repo=ChatRepository(db),
        stats_repo=StatsRepository(db)
    )

router = APIRouter(prefix="/chat", tags=["chat"])

@router.post("/conversations", response_model=ConversationResponse)
def create_conversation(
    workspace_id: Optional[int] = None,
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    return chat_service.create_conversation(current_user.id, workspace_id=workspace_id)

@router.get("/conversations", response_model=List[ConversationResponse])
def list_conversations(
    workspace_id: Optional[int] = None,
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    return chat_service.get_user_conversations(current_user.id, workspace_id=workspace_id)

@router.get("/conversations/{conversation_id}", response_model=ConversationResponse)
def get_conversation(
    conversation_id: str,
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    return chat_service.get_conversation_with_messages(conversation_id, current_user.id)

@router.put("/conversations/{conversation_id}", response_model=ConversationResponse)
def rename_conversation(
    conversation_id: str,
    rename_in: ConversationRenameRequest,
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    return chat_service.rename_conversation(conversation_id, rename_in.title, current_user.id)

@router.delete("/conversations/{conversation_id}")
def delete_conversation(
    conversation_id: str,
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    success = chat_service.delete_conversation(conversation_id, current_user.id)
    return {"success": success}

@router.post("/conversations/{conversation_id}/query")
async def query_conversation(
    request: Request,
    conversation_id: str,
    query_in: MessageCreate,
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    ip_address = request.client.host if request.client else None
    
    # Retrieve conversation scoping to resolve workspace_id filter
    conv = chat_service.get_conversation_with_messages(conversation_id, current_user.id)
    workspace_id = conv.workspace_id

    # Returns streaming response for SSE
    async def event_generator():
        # Set a keep-alive comment if processing takes too long
        async for chunk in chat_service.ask_question_stream(
            conversation_id=conversation_id,
            user_id=current_user.id,
            question=query_in.content,
            mode=query_in.mode or "chat",
            workspace_id=workspace_id,
            document_id=query_in.document_id,
            ip_address=ip_address
        ):
            yield chunk

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )

@router.post("/messages/{message_id}/feedback", response_model=MessageResponse)
def message_feedback(
    message_id: int,
    feedback_in: FeedbackRequest,
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    return chat_service.save_feedback(
        message_id=message_id,
        rating=feedback_in.feedback_rating,
        text=feedback_in.feedback_text,
        user_id=current_user.id
    )

@router.post("/messages/{message_id}/bookmark", response_model=MessageResponse)
def message_bookmark(
    message_id: int,
    bookmark_in: BookmarkRequest,
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    return chat_service.save_bookmark(
        message_id=message_id,
        bookmarked=bookmark_in.bookmarked,
        user_id=current_user.id
    )

@router.get("/bookmarks", response_model=List[MessageResponse])
def get_bookmarks(
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    return chat_service.get_bookmarks(current_user.id)

@router.get("/messages/{message_id}/export")
def export_message(
    message_id: int,
    format: str = Query(..., pattern="^(pdf|markdown|word)$"),
    chat_service: ChatService = Depends(get_chat_service),
    current_user: User = Depends(get_current_user)
):
    from app.models.chat import ChatMessage
    # Fetch message
    msg = chat_service.db.query(ChatMessage).filter(ChatMessage.id == message_id).first()
    if not msg:
        raise HTTPException(status_code=404, detail="Message not found")
    
    # Check ownership
    conv = chat_service.chat_repo.get_conversation(msg.conversation_id)
    if not conv or conv.user_id != current_user.id:
        raise HTTPException(status_code=403, detail="Access denied")
        
    content = msg.content
    
    if format == "markdown":
        import io
        stream = io.BytesIO(content.encode("utf-8"))
        return StreamingResponse(
            stream,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=message_{message_id}.md"}
        )
    elif format == "word":
        import io
        from docx import Document as DocxDocument
        doc = DocxDocument()
        doc.add_heading("Knowledge Platform Report", level=0)
        # Parse basic markdown-like syntax
        paragraphs = content.split("\n\n")
        for p in paragraphs:
            p_clean = p.strip()
            if p_clean.startswith("# "):
                doc.add_heading(p_clean.replace("# ", "").strip(), level=1)
            elif p_clean.startswith("## "):
                doc.add_heading(p_clean.replace("## ", "").strip(), level=2)
            else:
                doc.add_paragraph(p_clean)
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        return StreamingResponse(
            stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=message_{message_id}.docx"}
        )
    elif format == "pdf":
        import io
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        
        pdf_stream = io.BytesIO()
        doc = SimpleDocTemplate(
            pdf_stream,
            pagesize=letter,
            rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54
        )
        
        styles = getSampleStyleSheet()
        # Custom styles for a premium look
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=20,
            textColor=colors.HexColor('#1e3a8a'),
            spaceAfter=15
        )
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'),
            spaceAfter=8
        )
        h2_style = ParagraphStyle(
            'H2Style',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=13,
            textColor=colors.HexColor('#2563eb'),
            spaceAfter=10,
            spaceBefore=10
        )
        
        story = []
        story.append(Paragraph("Enterprise Knowledge Platform - AI Report", title_style))
        story.append(Spacer(1, 10))
        
        paragraphs = content.split("\n\n")
        for p in paragraphs:
            p_text = p.strip()
            if not p_text:
                continue
            if p_text.startswith("# "):
                story.append(Paragraph(p_text.replace("# ", "").strip(), h2_style))
            elif p_text.startswith("## "):
                story.append(Paragraph(p_text.replace("## ", "").strip(), h2_style))
            elif p_text.startswith("- ") or p_text.startswith("* "):
                # bullet list
                items = p_text.split("\n")
                for item in items:
                    item_text = item.strip().lstrip("-*").strip()
                    if item_text:
                        story.append(Paragraph(f"&bull; {item_text}", body_style))
            else:
                # normal text
                story.append(Paragraph(p_text.replace("\n", "<br/>"), body_style))
                
        doc.build(story)
        pdf_stream.seek(0)
        return StreamingResponse(
            pdf_stream,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=message_{message_id}.pdf"}
        )
