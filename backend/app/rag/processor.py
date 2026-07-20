import os
import csv
import pandas as pd
from typing import List, Dict, Any
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter

class DocumentProcessor:
    @staticmethod
    def parse_file(file_path: str, file_type: str) -> str:
        """
        Parses a file based on its extension/type and extracts clean text content.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_type.lower().strip(".")
        
        if file_type == "pdf":
            return DocumentProcessor._parse_pdf(file_path)
        elif file_type == "docx":
            return DocumentProcessor._parse_docx(file_path)
        elif file_type == "csv":
            return DocumentProcessor._parse_csv(file_path)
        elif file_type in ["txt", "md", "markdown"]:
            return DocumentProcessor._parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        text = []
        reader = PdfReader(file_path)
        for idx, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return "\n\n".join(text)

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        doc = DocxDocument(file_path)
        text = [paragraph.text for paragraph in doc.paragraphs]
        return "\n".join(text)

    @staticmethod
    def _parse_csv(file_path: str) -> str:
        df = pd.read_csv(file_path)
        rows_text = []
        for idx, row in df.iterrows():
            row_parts = [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
            rows_text.append(f"Row {idx + 1}: {', '.join(row_parts)}")
        return "\n".join(rows_text)

    @staticmethod
    def _parse_text(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """
        Splits text into chunks using recursive character splitting.
        """
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        return splitter.split_text(text)

    @staticmethod
    def parse_file_to_pages(file_path: str, file_type: str) -> List[Dict[str, Any]]:
        """
        Parses a file and returns a list of dictionaries, each containing:
        - "text": Clean extracted text
        - "page": The page number (if supported, else 1)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = file_type.lower().strip(".")
        
        if file_type == "pdf":
            pages = []
            reader = PdfReader(file_path)
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    pages.append({"text": page_text, "page": idx + 1})
            return pages
        elif file_type == "docx":
            doc = DocxDocument(file_path)
            text = [paragraph.text for paragraph in doc.paragraphs]
            return [{"text": "\n".join(text), "page": 1}]
        elif file_type == "csv":
            df = pd.read_csv(file_path)
            rows_text = []
            for idx, row in df.iterrows():
                row_parts = [f"{col}: {val}" for col, val in row.items() if pd.notna(val)]
                rows_text.append(f"Row {idx + 1}: {', '.join(row_parts)}")
            return [{"text": "\n".join(rows_text), "page": 1}]
        elif file_type in ["txt", "md", "markdown"]:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return [{"text": f.read(), "page": 1}]
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def chunk_pages(pages: List[Dict[str, Any]], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Dict[str, Any]]:
        """
        Splits pages of a document into chunks, keeping page-level metadata.
        """
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = []
        for page in pages:
            page_text = page["text"]
            page_num = page.get("page", 1)
            page_chunks = splitter.split_text(page_text)
            for pc in page_chunks:
                chunks.append({
                    "text": pc,
                    "metadata": {
                        "page_number": page_num
                    }
                })
        return chunks

    @staticmethod
    def chunk_pages_parent_child(
        pages: List[Dict[str, Any]], 
        parent_size: int = 2000, 
        child_size: int = 500, 
        child_overlap: int = 100
    ) -> List[Dict[str, Any]]:
        """
        Splits pages into large parent chunks and maps smaller child chunks to them.
        Saves parent text in child metadata for efficient retrieval.
        """
        parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=parent_size,
            chunk_overlap=int(parent_size * 0.15),
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=child_size,
            chunk_overlap=child_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        chunks = []
        for page in pages:
            page_text = page["text"]
            page_num = page.get("page", 1)
            
            # 1. Split into large Parent blocks
            parents = parent_splitter.split_text(page_text)
            for p_idx, parent_text in enumerate(parents):
                parent_id = f"page_{page_num}_parent_{p_idx}"
                
                # 2. Split Parent into smaller Child blocks
                children = child_splitter.split_text(parent_text)
                for c_idx, child_text in enumerate(children):
                    chunks.append({
                        "text": child_text,
                        "metadata": {
                            "page_number": page_num,
                            "parent_id": parent_id,
                            "parent_chunk_text": parent_text
                        }
                    })
        return chunks
